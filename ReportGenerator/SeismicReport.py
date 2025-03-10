from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, RGBColor, Pt, Cm

from .BasicGenerator import BasicGenerator, PageSize
from .DocParagraph import DocParagraph
from .DocPicture import DocPicture
from .SeismicReportTemplate import SRTemplate

class SeismicReport(BasicGenerator):
    def __init__(self):
        super().__init__()
        
        # 修改为A3图纸，横向，两栏
        self.change_paper_size(PageSize.A3_LANDSCAPE,2)
        self.change_paper_margin(32,25,32,25)
        
        # 格式统一修改
        self.body_style.paragraph_format.line_spacing = Pt(22)
        
    def creat_doc(self):
        self.__add_info()
        self.__add_seismic_chapter()
        
        
    def __add_info(self):
        model_name = "TestModel"
        par_context = SRTemplate.FIRST_INFO(model_name)
        paragraph = DocParagraph(par_context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)
        
    def __add_seismic_chapter(self):
        chapter_index = 8
        sub_index = 1
        self.__add_seismic_chapter_title(chapter_index)
        sub_index = self.__add_seismic_embedding(chapter_index,sub_index)
        
        
        
    def __add_seismic_chapter_title(self,chapter_index :int):
        
        yjk_version = "6.0.0"
        
        current_context = SRTemplate.SEISMIC_CHAPTER_TITLE
        par_context = DocParagraph(current_context.title(chapter_index))
        par_context.par_level  = 1
        self.add_title(par_context,12,6)
        paragraph_texts = current_context.paragraph(chapter_index,yjk_version)
        for context in paragraph_texts[:-1]:
            paragraph = DocParagraph(context)
            paragraph.style = self.body_style
            self.add_paragraph(paragraph)

        text = paragraph_texts[-1]
        paragraph = DocParagraph(text)
        paragraph.style = self.body_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.first_line_indent = 0
        self.add_paragraph(paragraph)
        
        figure_title = current_context.picture(chapter_index)
        paragraph = DocParagraph(figure_title)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
        
        
    def __add_seismic_embedding(self, chapter_index:int, sub_index:int):
        
        current_context = SRTemplate.SEISMIC_EMBEDDING
        par_context = DocParagraph(current_context.title(chapter_index,sub_index))
        par_context.par_level  = 2
        self.add_title(par_context,6,6)
        
        context = current_context.paragraph(chapter_index,sub_index)[0]
        paragraph = DocParagraph(context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)
        
        table_title = current_context.table(chapter_index,sub_index)
        paragraph = DocParagraph(table_title)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
        
        return sub_index + 1
    
    
    
    
    
    
    
    
    
    
    
    